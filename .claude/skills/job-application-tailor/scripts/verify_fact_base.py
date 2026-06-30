"""Verify that every technology/skill in cv_fact_base.json actually appears in the raw CV text.

Catches contamination where job offer keywords leak into the fact base during extraction.

- technologies and methodologies: STRICT — must appear literally in CV text. Failure = exit 1.
- skills: WARNING only — these are often reasonable abstractions of role descriptions
  (e.g. "Gestion d'équipe" from a "Development Manager" role). Flagged but not blocking.
- salient numeric metrics (carrying +, % or PB/TB/GB): STRICT — must still appear in the
  CV text (via factbase_consistency.find_metric_drift). Catches stale/fabricated figures
  like the "40+ → 100+ applications" drift that word-only checks missed. Failure = exit 1.

Exit code 0 = clean, exit code 1 = technology/methodology fabrication or metric drift found.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document

# Importable both as a bare script (preflight runs `python scripts/verify_fact_base.py`,
# putting the scripts dir on sys.path[0]) and as a package member (pytest puts the
# skill root on sys.path, so the module is `scripts.factbase_consistency`).
try:
    from factbase_consistency import find_metric_drift
except ImportError:  # pragma: no cover - exercised via the package-import path
    from scripts.factbase_consistency import find_metric_drift


def extract_cv_text(docx_path: Path) -> str:
    """Extract all text from a DOCX file (paragraphs + tables), lowercased."""
    doc = Document(str(docx_path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).lower()


# Known synonyms: fact base term -> also accept these in the CV text
_SYNONYMS: dict[str, list[str]] = {
    ".net core": [".net core", "dotnet core"],
    ".net": [".net", "dotnet"],
    "rest apis": ["rest", "api"],
    "gitlabci/cd": ["gitlab", "ci/cd"],
    "gitlab ci/cd": ["gitlab", "ci/cd"],
}


def _term_present(term: str, cv_text: str) -> bool:
    """Check if a term (or a known synonym) appears in the CV text."""
    term_lower = term.lower().strip()

    # Direct substring match
    if term_lower in cv_text:
        return True

    # Word-boundary regex match (handles e.g. "C#" which has special chars)
    escaped = re.escape(term_lower)
    if re.search(escaped, cv_text):
        return True

    # Check synonyms
    for syn in _SYNONYMS.get(term_lower, []):
        if syn in cv_text:
            return True

    return False


def verify(cv_path: Path, fact_base_path: Path) -> tuple[list[str], list[str]]:
    """Return (errors, warnings) for items not found in the raw CV text.

    errors: technologies/methodologies not in CV — blocks the pipeline.
    warnings: skills not in CV — logged but non-blocking.
    """
    cv_text = extract_cv_text(cv_path)

    with fact_base_path.open("r", encoding="utf-8") as f:
        fact_base = json.load(f)

    errors: list[str] = []
    warnings: list[str] = []

    # Strict: technologies and methodologies must appear literally
    for field in ("technologies", "methodologies"):
        for item in fact_base.get(field, []):
            if not _term_present(item, cv_text):
                errors.append(f"[{field}] {item}")

    # Warn-only: skills are often reasonable abstractions of role descriptions
    for item in fact_base.get("skills", []):
        if not _term_present(item, cv_text):
            warnings.append(f"[skills] {item}")

    # Strict: salient numeric metrics (carrying +, % or PB/TB/GB) must still
    # appear in the CV. verify_fact_base only ever checked tech/skill words, so a
    # changed figure (the "40+ → 100+ applications" incident) slipped through. A
    # drifted metric is a stale/fabricated number and blocks the pipeline, which
    # forces re-extraction on the cache-hit path (preflight calls this script).
    for token in find_metric_drift(cv_text, fact_base):
        errors.append(f"[metric] '{token}' in fact base but not in CV — stale/fabricated figure")

    return errors, warnings


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Verify CV fact base against raw CV text."
    )
    parser.add_argument("cv_path", help="Path to MASTER_CV.docx")
    parser.add_argument("fact_base_path", help="Path to cv_fact_base.json")
    args = parser.parse_args()

    errors, warnings = verify(Path(args.cv_path), Path(args.fact_base_path))

    if warnings:
        print(f"WARNING — {len(warnings)} skill(s) not found literally in CV (may be valid abstractions):")
        for item in warnings:
            print(f"  - {item}")
        print()

    if errors:
        print(f"FABRICATION DETECTED — {len(errors)} fact-base item(s) not found in CV:")
        for item in errors:
            print(f"  - {item}")
        print("\nRemove or re-extract these items from cv_fact_base.json before proceeding.")
        print("(A [metric] item means a number drifted from the CV — re-extract; do not refresh .cv_hash by hand.)")
        sys.exit(1)
    else:
        print("Verification OK — all technologies, methodologies, and metrics found in CV text.")


if __name__ == "__main__":
    main()
