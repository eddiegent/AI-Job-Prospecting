"""Build the fictional sample CV used by first-run onboarding (Phase 4).

Produces ``samples/MASTER_CV.example.docx`` — a neutral fictional CV
("Alex Dupont, Software Engineer") that uses the same style anchors as
``templates/cv_template_fr.docx`` so the extractor's heuristics still
match. This is **not** a Jinja template — it contains real literal
content a new user can open in Word and compare against their own CV.

Run once when the sample needs refreshing::

    python scripts/build_sample_cv.py
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

# Allow importing siblings when run directly (python scripts/build_sample_cv.py)
# as well as via module path (python -m scripts.build_sample_cv).
sys.path.insert(0, str(Path(__file__).resolve().parent))

from create_cv_template import (
    SKILL_ROOT,
    _create_styles,
    _para,
    _set_page,
)


SAMPLES_DIR = SKILL_ROOT / "samples"
OUT_PATH = SAMPLES_DIR / "MASTER_CV.example.docx"


def _header(doc: Document) -> None:
    _para(doc, "Alex Dupont", "NameStyle")
    _para(doc, "Software Engineer", "TaglineStyle")
    _para(doc, "Backend services, APIs, and data pipelines", "TaglineStyle")
    _para(
        doc,
        "alex.dupont@example.com  |  +33 6 00 00 00 00  |  "
        "linkedin.com/in/alexdupont-example  |  Lyon (69)",
        "ContactStyle",
    )


def _profile(doc: Document) -> None:
    _para(doc, "Professional Profile", "SectionStyle")
    _para(
        doc,
        "Software engineer with 8+ years designing and shipping backend "
        "services. Comfortable owning a feature from discovery through "
        "production support, and collaborating closely with product and "
        "data teams. Fluent in English and French working environments.",
    )


def _skills(doc: Document) -> None:
    _para(doc, "Skills", "SectionStyle")
    lines = [
        "Languages: Python, Go, TypeScript, SQL",
        "Frameworks: FastAPI, Django, Gin, React",
        "Data: PostgreSQL, Redis, Kafka, dbt, Airflow",
        "Cloud & DevOps: AWS (ECS, Lambda, S3, RDS), Docker, Terraform, GitHub Actions",
        "Practices: TDD, code review, pair programming, trunk-based development",
    ]
    for line in lines:
        _para(doc, line, "SkillLineStyle")


def _role(
    doc: Document,
    *,
    company_role_line: str,
    date_line: str,
    bullets: list[str],
) -> None:
    _para(doc, company_role_line, "RoleStyle")
    _para(doc, date_line, "SubtleStyle")
    for bullet in bullets:
        _para(doc, f"\u2022 {bullet}", "BulletStyle")


def _experience(doc: Document) -> None:
    _para(doc, "Professional Experience", "SectionStyle")

    _role(
        doc,
        company_role_line="Helios Analytics — Senior Backend Engineer",
        date_line="Mar 2022 – Present  |  Lyon, France",
        bullets=[
            "Led the migration of the billing service from a monolith to "
            "three FastAPI services, cutting p95 latency from 820 ms to 180 ms.",
            "Designed the event schema for a Kafka-based usage pipeline "
            "feeding downstream analytics (~12M events/day).",
            "Mentored two junior engineers through the team's TDD onboarding; "
            "both now own production services independently.",
        ],
    )

    _role(
        doc,
        company_role_line="Northbridge Software — Backend Engineer",
        date_line="Aug 2019 – Feb 2022  |  Paris, France",
        bullets=[
            "Built the public REST API for a B2B SaaS product used by 30+ "
            "enterprise clients, with OpenAPI-driven contract tests.",
            "Introduced database migrations discipline (Alembic) and "
            "removed a 6-month backlog of schema drift issues.",
            "On-call rotation for the order-management service; authored "
            "the runbook that reduced mean time to recovery by 40%.",
        ],
    )

    _role(
        doc,
        company_role_line="Northbridge Software — Software Engineer",
        date_line="Jun 2017 – Jul 2019  |  Paris, France",
        bullets=[
            "Implemented the reporting module in Django, including CSV "
            "export and scheduled email digests.",
            "Contributed to the React admin dashboard used by internal "
            "support staff.",
        ],
    )

    _role(
        doc,
        company_role_line="Earlier experience",
        date_line="",
        bullets=[
            "Junior developer and internship roles at small agencies in "
            "Lyon and Grenoble, working on PHP and early Python projects.",
        ],
    )


def _education(doc: Document) -> None:
    _para(doc, "Education", "SectionStyle")
    _para(
        doc,
        "2017 : INSA Lyon – Engineering degree (Diplôme d'ingénieur), "
        "Computer Science",
        "EducationStyle",
    )
    _para(
        doc,
        "2015 : IUT Grenoble – DUT Informatique (two-year technical degree)",
        "EducationStyle",
    )


def _languages(doc: Document) -> None:
    _para(doc, "Languages", "SectionStyle")
    _para(doc, "French (native)  |  English (fluent, C1)  |  Spanish (basic, A2)")


def build() -> Path:
    doc = Document()
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    _set_page(doc)
    _create_styles(doc)

    _header(doc)
    _profile(doc)
    _skills(doc)
    _experience(doc)
    _education(doc)
    _languages(doc)

    SAMPLES_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Created: {OUT_PATH}")
    return OUT_PATH


if __name__ == "__main__":
    build()
