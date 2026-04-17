"""Generate the CV .docx template(s) with styles and Jinja2 placeholder tags.

Run once (or whenever you want to refresh the template design):

    python scripts/create_cv_template.py

Outputs:
    templates/cv_template_fr.docx
    templates/cv_template_en.docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, Mm, RGBColor, Emu
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SKILL_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_DIR = SKILL_ROOT / "templates"

# ---------------------------------------------------------------------------
# Colours
# ---------------------------------------------------------------------------
BLUE = RGBColor(0x1F, 0x4E, 0x79)
DARK_GRAY = RGBColor(0x44, 0x44, 0x44)
SUBTLE_GRAY = RGBColor(0x60, 0x60, 0x60)
BLACK = RGBColor(0x00, 0x00, 0x00)

FONT = "Calibri"

# ---------------------------------------------------------------------------
# Section labels per language
# ---------------------------------------------------------------------------
LABELS = {
    "fr": {
        "profile": "Profil professionnel",
        "skills": "Comp\u00e9tences",
        "experience": "Exp\u00e9rience professionnelle",
        "education": "Formation",
        "languages": "Langues",
    },
    "en": {
        "profile": "Summary",
        "skills": "Skills",
        "experience": "Professional Experience",
        "education": "Education",
        "languages": "Languages",
    },
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _set_font(style, font_name: str) -> None:
    style.font.name = font_name
    rpr = style._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def _set_spacing(style, before_pt: float = 0, after_pt: float = 0) -> None:
    pf = style.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)


def _add_bottom_border(style, color_hex: str = "1F4E79", sz: str = "4") -> None:
    """Add a thin bottom border to a paragraph style (ATS-safe)."""
    pPr = style._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)          # half-points: 4 = 0.5pt
    bottom.set(qn("w:color"), color_hex)
    bottom.set(qn("w:space"), "1")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_keep_with_next(style) -> None:
    """Glue this paragraph to the next so Word won't orphan it at a page break."""
    pPr = style._element.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    pPr.append(keep)


def _para(doc: Document, text: str, style_name: str = "Normal") -> None:
    """Add a paragraph with the given text and style."""
    doc.add_paragraph(text, style=style_name)


# ---------------------------------------------------------------------------
# Style creation
# ---------------------------------------------------------------------------
def _create_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    _set_font(normal, FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.line_spacing = 1.08
    _set_spacing(normal, before_pt=0, after_pt=4)

    existing = {s.name for s in styles}

    # --- NameStyle ---
    if "NameStyle" not in existing:
        s = styles.add_style("NameStyle", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        _set_font(s, FONT)
        s.font.size = Pt(19)
        s.font.bold = True
        s.font.color.rgb = BLUE
        _set_spacing(s, before_pt=0, after_pt=2)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- TitleStyle (job title line, directly under the name) ---
    if "TitleStyle" not in existing:
        s = styles.add_style("TitleStyle", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        _set_font(s, FONT)
        s.font.size = Pt(16)
        s.font.bold = True
        s.font.color.rgb = BLUE
        _set_spacing(s, before_pt=0, after_pt=2)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- TaglineStyle (supporting line under the title, intentionally subtler) ---
    if "TaglineStyle" not in existing:
        s = styles.add_style("TaglineStyle", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        _set_font(s, FONT)
        s.font.size = Pt(10.5)
        s.font.italic = True
        s.font.color.rgb = SUBTLE_GRAY
        _set_spacing(s, before_pt=0, after_pt=4)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- ContactStyle ---
    if "ContactStyle" not in existing:
        s = styles.add_style("ContactStyle", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        _set_font(s, FONT)
        s.font.size = Pt(11)
        s.font.color.rgb = DARK_GRAY
        _set_spacing(s, before_pt=0, after_pt=2)
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- SectionStyle (with bottom border) ---
    if "SectionStyle" not in existing:
        s = styles.add_style("SectionStyle", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        _set_font(s, FONT)
        s.font.size = Pt(14)
        s.font.bold = True
        s.font.color.rgb = BLUE
        _set_spacing(s, before_pt=12, after_pt=4)
        _add_bottom_border(s)

    # --- RoleStyle ---
    if "RoleStyle" not in existing:
        s = styles.add_style("RoleStyle", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        _set_font(s, FONT)
        s.font.size = Pt(11)
        s.font.bold = True
        s.font.color.rgb = BLACK
        _set_spacing(s, before_pt=8, after_pt=0)
        _set_keep_with_next(s)

    # --- MetaStyle (Company | Location | Dates line, sits under the role) ---
    if "MetaStyle" not in existing:
        s = styles.add_style("MetaStyle", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        _set_font(s, FONT)
        s.font.size = Pt(10)
        s.font.italic = True
        s.font.color.rgb = SUBTLE_GRAY
        _set_spacing(s, before_pt=0, after_pt=3)
        _set_keep_with_next(s)

    # --- SubtleStyle (kept for backwards compatibility with older templates) ---
    if "SubtleStyle" not in existing:
        s = styles.add_style("SubtleStyle", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        _set_font(s, FONT)
        s.font.size = Pt(9.5)
        s.font.italic = True
        s.font.color.rgb = SUBTLE_GRAY
        _set_spacing(s, before_pt=0, after_pt=2)

    # --- SkillLineStyle ---
    if "SkillLineStyle" not in existing:
        s = styles.add_style("SkillLineStyle", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        _set_font(s, FONT)
        s.font.size = Pt(11)
        _set_spacing(s, before_pt=0, after_pt=2)

    # --- BulletStyle (list bullet without indentation issues) ---
    if "BulletStyle" not in existing:
        s = styles.add_style("BulletStyle", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        _set_font(s, FONT)
        s.font.size = Pt(11)
        _set_spacing(s, before_pt=0, after_pt=2)
        # Set left indent and hanging indent for bullet appearance
        pf = s.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.15)

    # --- EducationStyle ---
    if "EducationStyle" not in existing:
        s = styles.add_style("EducationStyle", WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = normal
        _set_font(s, FONT)
        s.font.size = Pt(11)
        _set_spacing(s, before_pt=0, after_pt=2)


# ---------------------------------------------------------------------------
# Template content with Jinja2 tags
# ---------------------------------------------------------------------------
def _add_template_content(doc: Document, labels: dict[str, str]) -> None:
    """Write the Jinja2-tagged placeholder content into the template."""

    # ── Header ──
    _para(doc, "{{candidate_name}}", "NameStyle")
    # Title (conditional — separate control paragraphs)
    _para(doc, "{%p if title %}")
    _para(doc, "{{title}}", "TitleStyle")
    _para(doc, "{%p endif %}")
    # Tagline (conditional, smaller than title)
    _para(doc, "{%p if tagline %}")
    _para(doc, "{{tagline}}", "TaglineStyle")
    _para(doc, "{%p endif %}")
    _para(doc, "{{r contact_rich}}", "ContactStyle")
    # Optional second contact line (long contact strings are split in half by the generator).
    _para(doc, "{%p if contact_rich_line2 %}")
    _para(doc, "{{r contact_rich_line2}}", "ContactStyle")
    _para(doc, "{%p endif %}")

    # ── Profile ──
    _para(doc, labels["profile"], "SectionStyle")
    _para(doc, "{%p for para in summary_paragraphs %}")
    _para(doc, "{{para}}")
    _para(doc, "{%p endfor %}")

    # ── Skills ──
    _para(doc, labels["skills"], "SectionStyle")
    _para(doc, "{%p for section in skills_sections %}")
    _para(doc, "{{r section.display}}", "SkillLineStyle")
    _para(doc, "{%p endfor %}")

    # ── Experience ──
    # Role (RoleStyle, bold) then "Company | Location | Dates" metadata line (MetaStyle).
    _para(doc, labels["experience"], "SectionStyle")
    _para(doc, "{%p for role in experience %}")
    _para(doc, "{{role.role_line}}", "RoleStyle")
    _para(doc, "{%p if role.metadata_line %}")
    _para(doc, "{{role.metadata_line}}", "MetaStyle")
    _para(doc, "{%p endif %}")
    _para(doc, "{%p for bullet in role.bullets %}")
    _para(doc, "\u2022 {{bullet}}", "BulletStyle")
    _para(doc, "{%p endfor %}")
    _para(doc, "{%p endfor %}")

    # ── Education ──
    _para(doc, "{%p if education %}")
    _para(doc, labels["education"], "SectionStyle")
    _para(doc, "{%p for line in education %}")
    _para(doc, "{{line}}", "EducationStyle")
    _para(doc, "{%p endfor %}")
    _para(doc, "{%p endif %}")

    # ── Languages ──
    _para(doc, "{%p if languages_line %}")
    _para(doc, labels["languages"], "SectionStyle")
    _para(doc, "{{languages_line}}")
    _para(doc, "{%p endif %}")


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------
def _set_page(doc: Document) -> None:
    section = doc.sections[0]
    # A4
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    # Margins
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def create_template(language: str = "fr") -> Path:
    labels = LABELS.get(language, LABELS["fr"])
    doc = Document()

    # Remove the default empty paragraph that python-docx adds
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    _set_page(doc)
    _create_styles(doc)
    _add_template_content(doc, labels)

    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    out_path = TEMPLATE_DIR / f"cv_template_{language}.docx"
    doc.save(str(out_path))
    print(f"Created: {out_path}")
    return out_path


def main() -> None:
    for lang in ("fr", "en"):
        create_template(lang)
    print("Done. Templates ready in:", TEMPLATE_DIR)


if __name__ == "__main__":
    main()
