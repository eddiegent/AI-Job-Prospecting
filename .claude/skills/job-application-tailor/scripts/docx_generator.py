from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxtpl import DocxTemplate, RichText
from jinja2 import Environment

from common import ensure_dir, load_yaml

SKILL_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE_DIR = SKILL_ROOT / "templates"


def _require(data: dict, field: str, context: str) -> Any:
    """Return data[field] or exit with a clear error."""
    if field not in data or not data[field]:
        print(f"ERROR in {context}: required field '{field}' is missing or empty.", file=sys.stderr)
        sys.exit(1)
    return data[field]


# ---------------------------------------------------------------------------
# CV template helpers
# ---------------------------------------------------------------------------

def _get_template_path(language: str) -> Path:
    """Return the CV template path for the given language, with fallback."""
    template = TEMPLATE_DIR / f"cv_template_{language}.docx"
    if template.exists():
        return template
    fallback = TEMPLATE_DIR / "cv_template_fr.docx"
    if fallback.exists():
        return fallback
    raise FileNotFoundError(
        f"CV template not found. Expected: {template}\n"
        f"Run 'python scripts/create_cv_template.py' to generate it."
    )


def _split_contact_lines(contact_line: str) -> list[str]:
    """Split a long contact line into readable halves at pipe separators.

    One line is kept when there are <= 3 pipe-separated items; four or more
    items split in half so the contact block never wraps mid-item on page.
    """
    if not contact_line:
        return [""]
    items = [p.strip() for p in contact_line.split("|") if p.strip()]
    if len(items) <= 3:
        return [" | ".join(items)]
    mid = (len(items) + 1) // 2
    return [" | ".join(items[:mid]), " | ".join(items[mid:])]


def _build_contact_richtext(tpl: DocxTemplate, contact_line: str) -> RichText:
    """Parse the contact line and create RichText with clickable hyperlinks."""
    rt = RichText()
    if not contact_line:
        return rt

    email_re = re.compile(r'[\w.+-]+@[\w-]+\.[\w.-]+')
    url_re = re.compile(r'(?:https?://)?(?:www\.)?linkedin\.com/\S+|(?:https?://)\S+')

    parts = re.split(r'(\s*\|\s*)', contact_line)
    for part in parts:
        email_match = email_re.search(part)
        url_match = url_re.search(part)
        if email_match:
            email = email_match.group()
            prefix = part[:email_match.start()]
            suffix = part[email_match.end():]
            if prefix:
                rt.add(prefix)
            rt.add(email, url_id=tpl.build_url_id(f"mailto:{email}"),
                   color="0563C1", underline=True)
            if suffix:
                rt.add(suffix)
        elif url_match:
            url = url_match.group()
            href = url if url.startswith("http") else f"https://{url}"
            prefix = part[:url_match.start()]
            suffix = part[url_match.end():]
            if prefix:
                rt.add(prefix)
            rt.add(url, url_id=tpl.build_url_id(href),
                   color="0563C1", underline=True)
            if suffix:
                rt.add(suffix)
        else:
            rt.add(part)
    return rt


def _build_skill_richtext(heading: str, items: list[str]) -> RichText:
    """Create a RichText with bold heading and normal-weight items."""
    rt = RichText()
    rt.add(f"{heading} : ", bold=True)
    rt.add(", ".join(items))
    return rt


def _load_section_labels(language: str) -> dict[str, str]:
    lang_file = SKILL_ROOT / "config" / "languages.yaml"
    if lang_file.exists():
        langs = load_yaml(lang_file)
        lang_data = langs.get(language, langs.get("fr", {}))
        return lang_data.get("cv_sections", {})
    return {}


_DEFAULT_LABELS = {
    "profile": "Profil professionnel",
    "skills": "Comp\u00e9tences",
    "experience": "Exp\u00e9rience professionnelle",
    "education": "Formation",
    "languages": "Langues",
}


# ---------------------------------------------------------------------------
# CV generation (template-based)
# ---------------------------------------------------------------------------

def generate_cv_docx(output_path: Path, cv_data: dict[str, Any], settings: dict[str, Any], language: str = "fr") -> Path:
    _require(cv_data, "candidate_name", "CV DOCX")
    _require(cv_data, "experience", "CV DOCX")

    ensure_dir(output_path.parent)
    labels = {**_DEFAULT_LABELS, **_load_section_labels(language)}
    template_path = _get_template_path(language)
    tpl = DocxTemplate(template_path)

    # Build RichText for the contact line(s). Long contact lines are split in
    # half so they don't wrap awkwardly mid-item at the page edge.
    contact_parts = _split_contact_lines(cv_data.get("contact_line", ""))
    contact_rich = _build_contact_richtext(tpl, contact_parts[0])
    contact_rich_line2 = _build_contact_richtext(tpl, contact_parts[1]) if len(contact_parts) > 1 else None

    # Build RichText for each skill section (bold heading + normal items)
    skills = []
    for section in cv_data.get("skills_sections", []):
        skills.append({
            "display": _build_skill_richtext(section["heading"], section["items"]),
        })

    context = {
        "candidate_name": cv_data.get("candidate_name", ""),
        "title": cv_data.get("title", ""),
        "tagline": cv_data.get("tagline", ""),
        "contact_rich": contact_rich,
        "contact_rich_line2": contact_rich_line2,
        "labels": labels,
        "summary_paragraphs": cv_data.get("summary_paragraphs", []),
        "skills_sections": skills,
        "experience": cv_data.get("experience", []),
        "education": cv_data.get("education", []),
        "languages_line": ", ".join(cv_data.get("languages", [])),
    }

    # Autoescape ensures `&`, `<`, `>` in values don't corrupt the rendered XML.
    tpl.render(context, jinja_env=Environment(autoescape=True))
    tpl.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Letter generation (unchanged — still uses python-docx directly)
# ---------------------------------------------------------------------------

def _set_font(style, font_name: str) -> None:
    """Set font name including East Asia fallback."""
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_default_font_letter(document: Document, font_name: str, body_size_pt: int) -> None:
    styles = document.styles
    normal = styles["Normal"]
    _set_font(normal, font_name)
    normal.font.size = Pt(body_size_pt)


def _set_margins(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def generate_letter_docx(output_path: Path, letter_data: dict[str, Any], settings: dict[str, Any]) -> Path:
    _require(letter_data, "paragraphs", "Letter DOCX")
    _require(letter_data, "name", "Letter DOCX")
    _require(settings, "formatting", "settings")

    ensure_dir(output_path.parent)
    formatting = settings["formatting"]
    doc = Document()
    _set_margins(doc)
    _set_default_font_letter(doc, formatting["font_name"], formatting["body_font_size_pt"])

    # Sender block (top-left)
    if letter_data.get("sender_name") or letter_data.get("sender_address"):
        if letter_data.get("sender_name"):
            p = doc.add_paragraph()
            r = p.add_run(letter_data["sender_name"])
            r.bold = True
        for line in letter_data.get("sender_address", []):
            doc.add_paragraph(line)
        doc.add_paragraph("")  # spacing

    # Recipient block
    if letter_data.get("recipient_name") or letter_data.get("recipient_address"):
        if letter_data.get("recipient_name"):
            p = doc.add_paragraph()
            r = p.add_run(letter_data["recipient_name"])
            r.bold = True
        for line in letter_data.get("recipient_address", []):
            doc.add_paragraph(line)
        doc.add_paragraph("")  # spacing

    if letter_data.get("date_line"):
        doc.add_paragraph(letter_data["date_line"])
    if letter_data.get("subject_line"):
        p = doc.add_paragraph()
        r = p.add_run(letter_data["subject_line"])
        r.bold = True
    if letter_data.get("greeting"):
        doc.add_paragraph(letter_data["greeting"])
    for para in letter_data.get("paragraphs", []):
        doc.add_paragraph(para)
    signoff = letter_data.get("signoff", "")
    name = letter_data.get("name", "")
    if signoff:
        doc.add_paragraph(signoff)
    if name:
        doc.add_paragraph(name)

    doc.save(str(output_path))
    return output_path
