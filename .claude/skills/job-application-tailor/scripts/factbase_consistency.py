"""Guardrail: detect when a cached cv_fact_base.json has drifted from the master CV.

**Canonical implementation.** This module is the single source of truth for the
fact-base consistency check. The repo-root ``tools/factbase_consistency.py`` is a
thin shim that re-exports everything here so the standalone CLI keeps working
without duplicating logic that could drift.

No skill-internal imports (only ``docx`` + stdlib), so it is importable both as a
bare module (``import factbase_consistency`` when the scripts dir is on
``sys.path``, e.g. when a sibling script runs directly) and as a package member
(``from scripts.factbase_consistency import check`` under pytest / preflight).

    python tools/factbase_consistency.py resources/MASTER_CV.docx resources/cv_fact_base.json --check-hash

Background
----------
The CV cache is keyed on the CV's SHA-256 hash. When the CV changes, the hash
no longer matches and preflight returns ``cache_stale`` — the signal to
RE-EXTRACT the fact base. The failure this guards against: reusing a stale fact
base and refreshing ``.cv_hash`` without re-extracting, which silently ships
outdated facts (the real incident: CV updated "40+ → 100+ applications" while
the fact base still said "40+").

``verify_fact_base.py`` historically only checked technologies/methodologies
appear in the CV — it never compared **numeric metrics**, so a changed number
slipped through. This module supplies the metric-drift check that closed that
hole; it is now wired into ``verify_fact_base.py``, ``common.save_cv_fact_base()``,
and ``preflight`` (see CLAUDE.md). Two deterministic checks:

1. **Hash freshness** (``--check-hash``) — stored ``.cv_hash`` must match the CV.
2. **Metric drift** — every *salient* numeric metric in the fact base (numbers
   carrying ``+``, ``%`` or a data-size unit PB/TB/GB) must still appear in the
   CV text. Bare integers and 4-digit years are ignored to stay
   false-positive-free.

Exit 0 = consistent, exit 1 = drift / stale hash.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document


def file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def extract_cv_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).lower()


def _normalize(text: str) -> str:
    """Lowercase, fold odd spaces, and remove whitespace *between digits* so
    '17 000+' and '17000+' compare equal."""
    text = unicodedata.normalize("NFKC", text).lower()
    text = text.replace(" ", " ").replace(" ", " ")
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"(?<=\d)\s+(?=\d)", "", text)
    return text


# Salient metric = a number carrying +, % or a data-size unit.
_METRIC_RE = re.compile(r"\d[\d   .,]*\s*(?:\+|%|\bpb\b|\btb\b|\bgb\b)", re.IGNORECASE)

_METRIC_FIELDS = ("summary",)


def _iter_metric_strings(fact_base: dict) -> list[str]:
    out: list[str] = []
    for field in _METRIC_FIELDS:
        v = fact_base.get(field)
        if isinstance(v, str):
            out.append(v)
    for exp in fact_base.get("experience", []):
        if isinstance(exp, dict):
            for key in ("details", "metrics"):
                out.extend(s for s in exp.get(key, []) if isinstance(s, str))
    for proj in fact_base.get("projects", []):
        if isinstance(proj, dict):
            out.extend(s for s in proj.get("bullets", []) if isinstance(s, str))
    return out


def find_metric_drift(cv_text: str, fact_base: dict) -> list[str]:
    """Return fact-base metrics absent from the CV text (stale/fabricated)."""
    cv_norm = _normalize(cv_text)
    drift: list[str] = []
    seen: set[str] = set()
    for s in _iter_metric_strings(fact_base):
        for raw in _METRIC_RE.findall(s):
            token = _normalize(raw)
            if token in seen:
                continue
            seen.add(token)
            if token not in cv_norm:
                drift.append(raw.strip())
    return drift


def check(cv_path: Path, fact_base_path: Path, check_hash: bool = False) -> tuple[list[str], list[str]]:
    errors: list[str] = []
    warnings: list[str] = []
    fact_base = json.loads(fact_base_path.read_text(encoding="utf-8"))
    cv_text = extract_cv_text(cv_path)

    for token in find_metric_drift(cv_text, fact_base):
        errors.append(f"[metric] '{token}' is in the fact base but not in the CV — stale or fabricated figure")

    if check_hash:
        hash_file = cv_path.parent / ".cv_hash"
        if hash_file.exists():
            if hash_file.read_text(encoding="utf-8").strip() != file_sha256(cv_path):
                errors.append("[hash] .cv_hash does not match the current CV — cache is stale, re-extract the fact base")
        else:
            warnings.append("[hash] no .cv_hash found next to the CV")
    return errors, warnings


def main() -> None:
    ap = argparse.ArgumentParser(description="Detect drift between a CV fact base and the master CV.")
    ap.add_argument("cv_path")
    ap.add_argument("fact_base_path")
    ap.add_argument("--check-hash", action="store_true")
    args = ap.parse_args()

    errors, warnings = check(Path(args.cv_path), Path(args.fact_base_path), args.check_hash)
    for w in warnings:
        print(f"WARNING — {w}")
    if errors:
        print(f"DRIFT DETECTED — {len(errors)} issue(s):")
        for e in errors:
            print(f"  - {e}")
        print("\nThe fact base is out of sync with the CV. Re-extract it (Steps 1-2.5); do not refresh .cv_hash by hand.")
        sys.exit(1)
    print("Consistency OK — fact base metrics match the CV" + (" and .cv_hash is fresh." if args.check_hash else "."))


if __name__ == "__main__":
    main()
