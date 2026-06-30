"""Guard: metric drift between cv_fact_base.json and the master CV must block.

Covers the real incident — the CV was updated "40+ → 100+ applications" while the
cached fact base still said "40+", and verify_fact_base only checked tech/skill
words, never numbers, so the stale figure shipped. These tests prove the closed
hole:

- ``verify_fact_base.verify()`` returns a blocking ``[metric]`` error on a drifted
  number (and none when consistent).
- ``common.save_cv_fact_base()`` raises rather than refreshing ``.cv_hash`` on a
  drifted fact base, so a stale cache can never be re-blessed.
"""
from __future__ import annotations

import json
from pathlib import Path

import pytest
from docx import Document

from scripts.common import save_cv_fact_base  # via conftest sys.path
from scripts.verify_fact_base import verify


def _make_cv(path: Path, *, applications: str = "100+") -> Path:
    """A minimal DOCX whose text carries two salient metrics: a configurable
    applications figure and a fixed '17 000+' backups figure."""
    doc = Document()
    doc.add_paragraph("Edward Gent — Ingénieur Logiciel Senior C# / .NET")
    doc.add_paragraph(
        f"Plus de {applications} candidatures envoyées en 2026. "
        "17 000+ sauvegardes quotidiennes pour plusieurs téraoctets de données."
    )
    doc.save(str(path))
    return path


def _fact_base(applications: str) -> dict:
    return {
        "summary": (
            f"Ingénieur logiciel senior ayant envoyé {applications} candidatures, "
            "avec 17 000+ sauvegardes quotidiennes en production."
        ),
        "technologies": ["C#", ".NET"],
        "methodologies": [],
        "skills": [],
        "experience": [],
    }


def test_consistent_metric_passes_verify(tmp_path):
    cv = _make_cv(tmp_path / "MASTER_CV.docx", applications="100+")
    fb = tmp_path / "cv_fact_base.json"
    fb.write_text(json.dumps(_fact_base("100+"), ensure_ascii=False), encoding="utf-8")
    errors, _ = verify(cv, fb)
    assert errors == [], errors


def test_drifted_metric_fails_verify(tmp_path):
    # CV says 100+, the cached fact base still says 40+ — the incident.
    cv = _make_cv(tmp_path / "MASTER_CV.docx", applications="100+")
    fb = tmp_path / "cv_fact_base.json"
    fb.write_text(json.dumps(_fact_base("40+"), ensure_ascii=False), encoding="utf-8")
    errors, _ = verify(cv, fb)
    assert any(e.startswith("[metric]") and "40+" in e for e in errors), errors


def test_save_refuses_drifted_fact_base(tmp_path):
    resources = tmp_path / "resources"
    prep = tmp_path / "_prep"
    resources.mkdir()
    prep.mkdir()
    cv = _make_cv(resources / "MASTER_CV.docx", applications="100+")
    (prep / "cv_fact_base.json").write_text(
        json.dumps(_fact_base("40+"), ensure_ascii=False), encoding="utf-8"
    )
    with pytest.raises(RuntimeError, match="inconsistent with the CV"):
        save_cv_fact_base(cv, prep)
    # The hash must NOT have been written — refreshing it on a stale fact base
    # is exactly the operation the guard exists to block.
    assert not (resources / ".cv_hash").exists()


def test_save_accepts_consistent_fact_base(tmp_path):
    resources = tmp_path / "resources"
    prep = tmp_path / "_prep"
    resources.mkdir()
    prep.mkdir()
    cv = _make_cv(resources / "MASTER_CV.docx", applications="100+")
    (prep / "cv_fact_base.json").write_text(
        json.dumps(_fact_base("100+"), ensure_ascii=False), encoding="utf-8"
    )
    save_cv_fact_base(cv, prep)
    assert (resources / ".cv_hash").exists()
    assert (resources / "cv_fact_base.json").exists()
